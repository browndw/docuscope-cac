# Copyright (C) 2024 David West Brown

# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at

#     http://www.apache.org/licenses/LICENSE-2.0

# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
import pathlib

import docx
from docx.shared import RGBColor
from docx.shared import Pt
import gzip
import glob
from io import BytesIO
import pickle
import streamlit as st
import pandas as pd
import polars as pl
import random
import zipfile
import xlsxwriter

HERE = pathlib.Path(__file__).parents[1].resolve()
CORPUS_DIR = HERE.joinpath("_corpora")
TEMP_DIR = HERE.joinpath("_temp")
OPTIONS = str(HERE.joinpath("options.toml"))
IMPORTS = str(HERE.joinpath("utilities/handlers_imports.py"))

# Functions for handling states and files.
# Handling can be done either by storing temporary files locally or by storing data in session memory.
# Thus, functions are written in matching pairs, depending on how data is to be stored.
# These functions can be imported either from handlers_server.py or from handlers_local.py as _handlers
# When imported, then, a given function call is the same, taking the same arguments.
# And session data can be imported and stored as needed.

# Handling is set by the enable_save boolean value in options.toml

# Initialize session states.
# For local file handling, generate a temp folder to be deleted on close.


def generate_temp(states, session_id):
	if session_id not in st.session_state:
		st.session_state[session_id] = {}
	for key, value in states:
		if key not in st.session_state[session_id]:
			st.session_state[session_id][key] = value

def init_session(session_id):
	session = {}
	session['has_target'] = False
	session['target_db'] = ''
	session['has_meta'] = False
	session['has_reference'] = False
	session['reference_db'] = ''
	session['freq_table'] = False
	session['tags_table'] = False
	session['keyness_table'] = False
	session['ngrams'] = False
	session['kwic'] = False
	session['keyness_parts'] = False
	session['dtm'] = False
	session['pca'] = False
	session['collocations'] = False
	session['doc'] = False

	df = pl.from_dict(session)
	st.session_state[session_id]["session"] = df

# Functions for managing session values.

def update_session(key, value, session_id):
	session = st.session_state[session_id]["session"]
	session = session.to_dict(as_series=False)
	session[key] = value
	df = pl.from_dict(session)
	st.session_state[session_id]["session"] = df

# Functions for storing and managing corpus metadata

def init_metadata_target(session_id):
	df = st.session_state[session_id]["target"]["ds_tokens"]
	tags_to_check = df.get_column("ds_tag").to_list()
	tags = ['Actors', 'Organization', 'Planning', 'Sentiment', 'Signposting', 'Stance']
	if any(tag in item for item in tags_to_check for tag in tags):
		model = 'Common Dictionary'
	else:
		model = 'Large Dictionary'
	ds_tags = df.get_column("ds_tag").unique().to_list()
	tags_pos = df.get_column("pos_tag").unique().to_list()
	if "Untagged" in ds_tags:
		ds_tags.remove("Untagged")
	if "Y" in tags_pos:
		tags_pos.remove("Y")
	temp_metadata_target = {}
	temp_metadata_target['tokens_pos'] = df.group_by(["doc_id", "pos_id", "pos_tag"]).agg(pl.col("token").str.concat("")).filter(pl.col("pos_tag") != "Y").height
	temp_metadata_target['tokens_ds'] = df.group_by(["doc_id", "ds_id", "ds_tag"]).agg(pl.col("token").str.concat("")).filter(~(pl.col("token").str.contains("^[[[:punct:]] ]+$") & pl.col("ds_tag").str.contains("Untagged"))).height
	temp_metadata_target['ndocs'] = len(df.get_column("doc_id").unique().to_list())
	temp_metadata_target['model'] = model
	temp_metadata_target['docids'] = {'ids': sorted(df.get_column("doc_id").unique().to_list())}
	temp_metadata_target['tags_ds'] = {'tags': sorted(ds_tags)}
	temp_metadata_target['tags_pos'] = {'tags': sorted(tags_pos)}
	temp_metadata_target['doccats'] = {'cats': ''}
	temp_metadata_target['collocations'] = {'temp': ''}
	temp_metadata_target['keyness_parts'] = {'temp': ''}
	temp_metadata_target['variance'] = {'temp': ''}

	df = pl.from_dict(temp_metadata_target, strict=False)
	st.session_state[session_id]["metadata_target"] = df

def init_metadata_reference(session_id):
	df = st.session_state[session_id]["reference"]["ds_tokens"]
	tags_to_check = df.get_column("ds_tag").to_list()
	tags = ['Actors', 'Organization', 'Planning', 'Sentiment', 'Signposting', 'Stance']
	if any(tag in item for item in tags_to_check for tag in tags):
		model = 'Common Dictionary'
	else:
		model = 'Large Dictionary'
	ds_tags = df.get_column("ds_tag").unique().to_list()
	tags_pos = df.get_column("pos_tag").unique().to_list()
	if "Untagged" in ds_tags:
		ds_tags.remove("Untagged")
	if "Y" in tags_pos:
		tags_pos.remove("Y")
	temp_metadata_reference = {}
	temp_metadata_reference['tokens_pos'] = df.group_by(["doc_id", "pos_id", "pos_tag"]).agg(pl.col("token").str.concat("")).filter(pl.col("pos_tag") != "Y").height
	temp_metadata_reference['tokens_ds'] = df.group_by(["doc_id", "ds_id", "ds_tag"]).agg(pl.col("token").str.concat("")).filter(~(pl.col("token").str.contains("^[[[:punct:]] ]+$") & pl.col("ds_tag").str.contains("Untagged"))).height
	temp_metadata_reference['ndocs'] = len(df.get_column("doc_id").unique().to_list())
	temp_metadata_reference['model'] = model
	temp_metadata_reference['doccats'] = False
	temp_metadata_reference['docids'] = {'ids': sorted(df.get_column("doc_id").unique().to_list())}
	temp_metadata_reference['tags_ds'] = {'tags': sorted(ds_tags)}
	temp_metadata_reference['tags_pos'] = {'tags': sorted(tags_pos)}

	df = pl.from_dict(temp_metadata_reference, strict=False)
	st.session_state[session_id]["metadata_reference"] = df

def load_metadata(corpus_type, session_id):
	table_name = "metadata_" + corpus_type
	metadata = st.session_state[session_id][table_name]
	metadata = metadata.to_dict(as_series=False)
	return(metadata)

def update_metadata(corpus_type, key, value, session_id):
	table_name = "metadata_" + corpus_type
	metadata = st.session_state[session_id][table_name]
	metadata = metadata.to_dict(as_series=False)
	if key == "doccats":
		metadata['doccats'] = {'cats': [value]}
	elif key == "collocations":
		metadata['collocations'] = {'temp': [value]}
	elif key == "keyness_parts":
		metadata['keyness_parts'] = {'temp': [value]}
	elif key == "variance":
		metadata['variance'] = {'temp': [value]}
	else:
		metadata[key] = value
	df = pl.from_dict(metadata, strict=False)
	st.session_state[session_id][table_name] = df

# Functions for handling corpora

def load_corpus_internal(db_path, session_id, corpus_type='target'):
	if corpus_type not in st.session_state[session_id]:
		st.session_state[session_id][corpus_type] = {}
	files_list = glob.glob(os.path.join(db_path, '*.gz'))
	random.shuffle(files_list)
	data = {}
	for file in files_list:
		try:
			with gzip.open(file, 'rb') as f:
				data[str(os.path.basename(file)).removesuffix(".gz")] = pickle.load(f)
		except:
			pass	
	if len(data) != 7:
		random.shuffle(files_list)
		data = {}
		for file in files_list:
			try:
				with gzip.open(file, 'rb') as f:
					data[str(os.path.basename(file)).removesuffix(".gz")] = pickle.load(f)
			except:
				pass	
	else:
		for key, value in data.items():
			if key not in st.session_state[session_id][corpus_type]:
				st.session_state[session_id][corpus_type][key] = {}
			st.session_state[session_id][corpus_type][key] = value

def load_corpus_new(ds_tokens,
					dtm_ds,
					dtm_pos,
					ft_ds,
					ft_pos,
					tt_ds,
					tt_pos,
					session_id, 
					corpus_type='target'):

	if corpus_type not in st.session_state[session_id]:
		st.session_state[session_id][corpus_type] = {}
	if "ds_tokens" not in st.session_state[session_id][corpus_type]:
		st.session_state[session_id][corpus_type]["ds_tokens"] = {}
	st.session_state[session_id][corpus_type]["ds_tokens"] = ds_tokens
	if "dtm_ds" not in st.session_state[session_id][corpus_type]:
		st.session_state[session_id][corpus_type]["dtm_ds"] = {}
	st.session_state[session_id][corpus_type]["dtm_ds"] = dtm_ds
	if "dtm_pos" not in st.session_state[session_id][corpus_type]:
		st.session_state[session_id][corpus_type]["dtm_pos"] = {}
	st.session_state[session_id][corpus_type]["dtm_pos"] = dtm_pos
	if "ft_ds" not in st.session_state[session_id][corpus_type]:
		st.session_state[session_id][corpus_type]["ft_ds"] = {}
	st.session_state[session_id][corpus_type]["ft_ds"] = ft_ds	
	if "ft_pos" not in st.session_state[session_id][corpus_type]:
		st.session_state[session_id][corpus_type]["ft_pos"] = {}
	st.session_state[session_id][corpus_type]["ft_pos"] = ft_pos	
	if "tt_ds" not in st.session_state[session_id][corpus_type]:
		st.session_state[session_id][corpus_type]["tt_ds"] = {}
	st.session_state[session_id][corpus_type]["tt_ds"] = tt_ds	
	if "tt_pos" not in st.session_state[session_id][corpus_type]:
		st.session_state[session_id][corpus_type]["tt_pos"] = {}
	st.session_state[session_id][corpus_type]["tt_pos"] = tt_pos

def find_saved(model_type: str):
	SUB_DIR = CORPUS_DIR.joinpath(model_type)
	saved_paths = [ f.path for f in os.scandir(SUB_DIR) if f.is_dir() ]
	saved_names = [ f.name for f in os.scandir(SUB_DIR) if f.is_dir() ]
	saved_corpora = {saved_names[i]: saved_paths[i] for i in range(len(saved_names))}
	return(saved_corpora)

def find_saved_reference(target_model, target_path):
	# only allow comparisions of ELSEVIER to MICUSP
	target_base = os.path.splitext(os.path.basename(pathlib.Path(target_path)))[0]
	if "MICUSP" in target_base:
		corpus = "MICUSP"
	else:
		corpus = "ELSEVIER"
	model_type = ''.join(word[0] for word in target_model.lower().split())
	SUB_DIR = CORPUS_DIR.joinpath(model_type)
	saved_paths = [ f.path for f in os.scandir(SUB_DIR) if f.is_dir() ]
	saved_names = [ f.name for f in os.scandir(SUB_DIR) if f.is_dir() ]
	saved_corpora = {saved_names[i]: saved_paths[i] for i in range(len(saved_names))}
	saved_ref = {key:val for key, val in saved_corpora.items() if corpus not in key}
	return(saved_corpora, saved_ref)

# Functions for handling data tables

def convert_to_parquet(_df):
	_df = pl.from_arrow(_df).to_pandas()
	_df = _df.to_parquet()
	return _df

def convert_to_excel(_df):
	output = BytesIO()
	writer = pd.ExcelWriter(output, engine='xlsxwriter')
	_df.to_excel(writer, index=False, header=True)
	writer.close()
	processed_data = output.getvalue()
	return processed_data

def add_alt_chunk(doc: docx.Document, html: str):
    package = doc.part.package
    partname = package.next_partname('/word/altChunk%d.html')
    alt_part = docx.opc.part.Part(partname, 'text/html', html.encode(), package)
    r_id = doc.part.relate_to(alt_part, docx.opc.constants.RELATIONSHIP_TYPE.A_F_CHUNK)
    alt_chunk = docx.oxml.OxmlElement('w:altChunk')
    alt_chunk.set(docx.oxml.ns.qn('r:id'), r_id)
    doc.element.body.sectPr.addprevious(alt_chunk)

def convert_to_word(html_string, tag_html, doc_key, tag_counts):
	doc_html = html_string.split('</style>')
	style_sheet_str = doc_html[0] + '</style>'
	html_str = doc_html[1]
	doc_html = '<!DOCTYPE html><html><head>' + style_sheet_str + '</head><body>' + tag_html + '<br><br>' + html_str + '</body></html>'
	download_file = docx.Document()
	title = download_file.add_heading(doc_key)
	title.style.font.color.rgb = RGBColor(0, 0, 0)
	heading = download_file.add_heading('Table of tag frequencies:', 3)
	heading.style.font.color.rgb = RGBColor(0, 0, 0)
	#add counts table
	tag_counts['RF'] = tag_counts.RF.round(2)
	t = download_file.add_table(tag_counts.shape[0]+1, tag_counts.shape[1])
	# add the header rows.
	for j in range(tag_counts.shape[-1]):
		t.cell(0,j).text = tag_counts.columns[j]
	# add the rest of the data frame
	for i in range(tag_counts.shape[0]):
		for j in range(tag_counts.shape[-1]):
			t.cell(i+1,j).text = str(tag_counts.values[i,j])
	t.style = 'LightList'
	for row in t.rows:
		for cell in row.cells:
			paragraphs = cell.paragraphs
			for paragraph in paragraphs:
				for run in paragraph.runs:
					font = run.font
					font.size = Pt(10)
					font.name = "Arial"
	download_file.add_heading('Highlighted tags:', 3)
	#add html
	add_alt_chunk(download_file, doc_html)
	output = BytesIO()
	download_file.save(output)
	processed_data = output.getvalue()
	return processed_data

def convert_corpus_to_zip(session_id, corpus_type, file_type="parquet"):
	zip_buf = BytesIO()
	with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as file_zip:
		for table in st.session_state[session_id][corpus_type]:
			_df = st.session_state[session_id][corpus_type][table]
			if file_type == "parquet":
				_df = _df.to_pandas().to_parquet()
				file_zip.writestr(table + ".parquet", _df)
			else:
				_df = _df.to_pandas().to_csv()
				file_zip.writestr(table + ".csv", _df)
	processed_data = zip_buf.getvalue()
	return(processed_data)

def convert_to_zip(tok_pl, tagset):
	zip_buf = BytesIO()
	with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as file_zip:
		for id in tok_pl.get_column("doc_id").unique().to_list():
				if tagset == "pos":
					df = (
						tok_pl
						.filter(pl.col("doc_id") == id)
						.group_by(["pos_id", "pos_tag"], maintain_order = True)
						.agg(pl.col("token").str.concat(""))
						.with_columns(pl.col("token").str.strip_chars())
						.with_columns(pl.col("token").str.replace_all(" ", "_"))
						.with_columns(pl.when(pl.col("pos_tag") == "Y").then(pl.col("pos_tag").str.replace("Y", "", literal=True))
									.when(pl.col("pos_tag") == "FU").then(pl.col("pos_tag").str.replace("FU", "", literal=True))
									.otherwise(pl.col("pos_tag")))
						.with_columns(pl.concat_str(pl.col("token"), pl.lit("|"), pl.col("pos_tag")))
						.with_columns(pl.col("token").str.replace_all("\|$", ""))
						)
				else:
					df = (
						tok_pl
						.filter(pl.col("doc_id") == id)
						.group_by(["ds_id", "ds_tag"], maintain_order = True)
						.agg(pl.col("token").str.concat(""))
						.with_columns(pl.col("token").str.strip_chars())
						.with_columns(pl.col("token").str.replace_all(" ", "_"))
						.with_columns(pl.when(pl.col("ds_tag") == "Untagged")
									.then(pl.col("ds_tag").str.replace("Untagged", "", literal=True))
									.otherwise(pl.col("ds_tag")))
						.with_columns(pl.concat_str(pl.col("token"), pl.lit("|"), pl.col("ds_tag")))
						.with_columns(pl.col("token").str.replace_all("\|$", ""))
						)
				doc = " ".join(df.get_column("token").to_list())
				file_zip.writestr(id + "_tagged"+ ".txt", doc)
	processed_data = zip_buf.getvalue()
	return(processed_data)

# Functions for storing values associated with specific apps

def update_tags(html_state, session_id):
	_TAGS = f"tags_{session_id}"
	html_highlights = [' { background-color:#5fb7ca; }', ' { background-color:#e35be5; }', ' { background-color:#ffc701; }', ' { background-color:#fe5b05; }', ' { background-color:#cb7d60; }']
	if 'html_str' not in st.session_state[session_id]:
		st.session_state[session_id]['html_str'] = ''
	if _TAGS in st.session_state:
		tags = st.session_state[_TAGS]
		if len(tags)>5:
			tags = tags[:5]
			st.session_state[_TAGS] = tags
	else:
		tags = []
	tags = ['.' + x for x in tags]
	highlights = html_highlights[:len(tags)]
	style_str = [''.join(x) for x in zip(tags, highlights)]
	style_str = ''.join(style_str)
	style_sheet_str = '<style>' + style_str + '</style>'
	st.session_state[session_id]['html_str'] = style_sheet_str + html_state

# Convenience function called by widgets

def clear_plots(session_id):
	update_session('pca', False, session_id)
	_GRPA = f"grpa_{session_id}"
	_GRPB = f"grpb_{session_id}"
	if _GRPA in st.session_state.keys():
		st.session_state[_GRPA] = []
	if _GRPB in st.session_state.keys():
		st.session_state[_GRPB] = []

def persist(key: str, app_name: str, session_id):
	_PERSIST_STATE_KEY = f"{app_name}_PERSIST"
	if _PERSIST_STATE_KEY not in st.session_state[session_id].keys():
		st.session_state[session_id][_PERSIST_STATE_KEY] = {}
		st.session_state[session_id][_PERSIST_STATE_KEY][key] = None
    
	if key in st.session_state:
		st.session_state[session_id][_PERSIST_STATE_KEY][key] = st.session_state[key]
    	
	return key
	
def load_widget_state(app_name: str, session_id):
	_PERSIST_STATE_KEY = f"{app_name}_PERSIST"
	"""Load persistent widget state."""
	if _PERSIST_STATE_KEY in st.session_state[session_id]:
		for key in st.session_state[session_id][_PERSIST_STATE_KEY]:
			if st.session_state[session_id][_PERSIST_STATE_KEY][key] is not None:
				if key not in st.session_state:
					st.session_state[key] = st.session_state[session_id][_PERSIST_STATE_KEY][key]

#prevent categories from being chosen in both multiselect
def update_grpa(session_id):
	_GRPA = f"grpa_{session_id}"
	_GRPB = f"grpb_{session_id}"
	if _GRPA not in st.session_state.keys():
		st.session_state[_GRPA] = []
	if _GRPB not in st.session_state.keys():
		st.session_state[_GRPB] = []
	if len(list(set(st.session_state[_GRPA]) & set(st.session_state[_GRPB]))) > 0:
		item = list(set(st.session_state[_GRPA]) & set(st.session_state[_GRPB]))
		st.session_state[_GRPA] = list(set(list(st.session_state[_GRPA]))^set(item))

def update_grpb(session_id):
	_GRPA = f"grpa_{session_id}"
	_GRPB = f"grpb_{session_id}"
	if _GRPA not in st.session_state.keys():
		st.session_state[_GRPA] = []
	if _GRPB not in st.session_state.keys():
		st.session_state[_GRPB] = []
	if len(list(set(st.session_state[_GRPA]) & set(st.session_state[_GRPB]))) > 0:
		item = list(set(st.session_state[_GRPA]) & set(st.session_state[_GRPB]))
		st.session_state[_GRPB] = list(set(list(st.session_state[_GRPB]))^set(item))

#prevent categories from being chosen in both multiselect
def update_tar(session_id):
	_TAR = f"tar_{session_id}"
	_REF = f"ref_{session_id}"
	if _TAR not in st.session_state.keys():
		st.session_state[_TAR] = []
	if _REF not in st.session_state.keys():
		st.session_state[_REF] = []
	if len(list(set(st.session_state[_TAR]) & set(st.session_state[_REF]))) > 0:
		item = list(set(st.session_state[_TAR]) & set(st.session_state[_REF]))
		st.session_state[_TAR] = list(set(list(st.session_state[_TAR]))^set(item))
		
def update_ref(session_id):
	_REF = f"ref_{session_id}"
	_TAR = f"tar_{session_id}"
	if _TAR not in st.session_state.keys():
		st.session_state[_TAR] = []
	if _REF not in st.session_state.keys():
		st.session_state[_REF] = []
	if len(list(set(st.session_state[_TAR]) & set(st.session_state[_REF]))) > 0:
		item = list(set(st.session_state[_TAR]) & set(st.session_state[_REF]))
		st.session_state[_REF] = list(set(list(st.session_state[_REF]))^set(item))


